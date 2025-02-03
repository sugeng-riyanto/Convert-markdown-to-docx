import streamlit as st
import sqlite3
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from markdown_it import MarkdownIt
from bs4 import BeautifulSoup

# Connect to SQLite database (or create it if it doesn't exist)
conn = sqlite3.connect('markdown_files.db', check_same_thread=False)
c = conn.cursor()

# Create table if it doesn't exist
c.execute('''
    CREATE TABLE IF NOT EXISTS markdown_files (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT,
        content TEXT
    )
''')
conn.commit()

# Function to save Markdown file to the database
def save_to_database(filename, content):
    c.execute('''
        INSERT INTO markdown_files (filename, content)
        VALUES (?, ?)
    ''', (filename, content))
    conn.commit()

# Function to fetch all Markdown files from the database
def fetch_all_files():
    c.execute('SELECT id, filename FROM markdown_files')
    return c.fetchall()

# Function to fetch content of a specific Markdown file by ID
def fetch_file_content(file_id):
    c.execute('SELECT content FROM markdown_files WHERE id = ?', (file_id,))
    return c.fetchone()[0]

# Function to delete a Markdown file from the database
def delete_file_from_database(file_id):
    c.execute('DELETE FROM markdown_files WHERE id = ?', (file_id,))
    conn.commit()

# Function to convert Markdown to DOCX with proper formatting
def markdown_to_docx(md_content, output_filename):
    # Initialize Markdown parser
    md = MarkdownIt()

    # Convert Markdown to HTML
    html_content = md.render(md_content)

    # Parse HTML using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create DOCX document
    doc = Document()

    def process_element(element, parent_paragraph=None):
        """Recursively process HTML elements."""
        if element.name == 'h1':
            heading = doc.add_heading(level=1)
            run = heading.add_run(element.text)
            run.bold = True
            run.font.size = Pt(16)

        elif element.name == 'h2':
            heading = doc.add_heading(level=2)
            run = heading.add_run(element.text)
            run.bold = True
            run.font.size = Pt(14)

        elif element.name == 'h3':
            heading = doc.add_heading(level=3)
            run = heading.add_run(element.text)
            run.bold = True
            run.font.size = Pt(12)

        elif element.name == 'p':
            paragraph = parent_paragraph or doc.add_paragraph()
            for child in element.children:
                if isinstance(child, str):  # Plain text
                    run = paragraph.add_run(child)
                    run.font.size = Pt(11)
                elif child.name == 'strong':  # Bold text
                    run = paragraph.add_run(child.text)
                    run.bold = True
                elif child.name == 'em':  # Italic text
                    run = paragraph.add_run(child.text)
                    run.italic = True
                elif child.name == 'a':  # Hyperlink
                    hyperlink = paragraph.add_hyperlink(child['href'])
                    hyperlink.text = child.text
                    hyperlink.style = "Hyperlink"
                elif child.name == 'img':  # Image
                    image_url = child.get('src', '')
                    try:
                        doc.add_picture(image_url, width=Inches(4))  # Adjust width as needed
                    except Exception as e:
                        doc.add_paragraph(f"Image could not be loaded: {image_url}")
                elif child.name == 'span' and 'class' in child.attrs and 'math' in child['class']:
                    # Handle LaTeX math equations
                    equation_text = child.text
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(equation_text)
                    run.font.size = Pt(11)
                    run.font.name = 'Cambria Math'

        elif element.name == 'table':
            # Handle tables
            table = doc.add_table(rows=0, cols=len(element.find_all('th') or element.find_all('td')))
            table.style = 'Table Grid'

            # Add header row if present
            header_row = element.find('tr')
            if header_row:
                cells = header_row.find_all(['th', 'td'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.text

            # Add data rows
            for row in element.find_all('tr')[1:]:
                cells = row.find_all(['th', 'td'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.text

    # Process all top-level elements
    for element in soup:
        process_element(element)

    # Save the DOCX file
    doc.save(output_filename)

# Sidebar for navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ["Upload Markdown", "View Markdown Files"])

# Upload Markdown Page
if page == "Upload Markdown":
    st.title("Upload Markdown File")

    # File uploader for Markdown file
    uploaded_file = st.file_uploader("Upload a Markdown file", type=["md"])

    if uploaded_file is not None:
        # Read the content of the uploaded Markdown file
        md_content = uploaded_file.read().decode("utf-8")

        # Display the Markdown content
        st.subheader("Uploaded Markdown Content:")
        st.text(md_content)

        # Save the Markdown content to the database
        if st.button("Save to Database"):
            save_to_database(uploaded_file.name, md_content)
            st.success(f"File '{uploaded_file.name}' saved to database!")

# View Markdown Files Page
elif page == "View Markdown Files":
    st.title("View Saved Markdown Files")

    # Fetch all saved Markdown files from the database
    files = fetch_all_files()

    if files:
        st.subheader("List of Saved Markdown Files:")

        # Dropdown to select a file for preview
        selected_file = st.selectbox("Select a Markdown file to preview", [filename for _, filename in files])

        # Get the selected file's ID and content
        selected_file_id = next(id for id, filename in files if filename == selected_file)
        selected_file_content = fetch_file_content(selected_file_id)

        # Show the full-page Markdown preview
        st.subheader(f"Preview: {selected_file}")
        st.markdown(selected_file_content, unsafe_allow_html=True)  # Render Markdown

        # Buttons for actions (Delete, Download as DOCX)
        col1, col2 = st.columns(2)

        with col1:
            # Button to delete the file
            if st.button(f"Delete {selected_file}", key=f"delete_{selected_file_id}"):
                delete_file_from_database(selected_file_id)
                st.experimental_rerun()  # Refresh the page after deletion

        with col2:
            # Button to download as DOCX
            if st.button(f"Download {selected_file} as DOCX", key=f"docx_{selected_file_id}"):
                docx_filename = f"{os.path.splitext(selected_file)[0]}.docx"
                markdown_to_docx(selected_file_content, docx_filename)

                with open(docx_filename, "rb") as docx_file:
                    st.download_button(
                        label=f"Download {selected_file} as DOCX",
                        data=docx_file,
                        file_name=docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Clean up temporary DOCX file
                os.remove(docx_filename)
    else:
        st.info("No Markdown files saved yet.")

# Close the database connection when the app is closed
conn.close()
